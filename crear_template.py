"""Script de utilidad para crear un template de memorandum de ejemplo."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def crear_template_memorandum():
    doc = Document()

    # Title
    titulo = doc.add_heading("MEMORANDUM", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Metadata table
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    filas = [
        ("PARA:", "<NombreEmpleado>"),
        ("ASUNTO:", "<Asunto>"),
        ("FECHA:", "<fecha>"),
        ("DOCUMENTO:", "<Documento>"),
    ]

    for i, (label, value) in enumerate(filas):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph("")

    # Body
    doc.add_paragraph(
        "Por medio del presente se le comunica que se le han autorizado <días> días "
        "de permiso a partir de la fecha indicada."
    )

    doc.add_paragraph("")
    doc.add_paragraph("Atentamente,")
    doc.add_paragraph("")
    doc.add_paragraph("_________________________")
    doc.add_paragraph("Firma autorizada")

    doc.save("templates/memorandum.docx")
    print("Template creado en templates/memorandum.docx")


if __name__ == "__main__":
    crear_template_memorandum()
