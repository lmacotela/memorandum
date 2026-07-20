from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from docx import Document
import io
import os
import uuid
import zipfile

from sapiens_client import obtener_token, obtener_memorandums
from word_service import generar_documento

app = FastAPI(title="Memorandum Word API", version="2.0.0")

TEMPLATES_DIR = "templates"
OUTPUT_DIR = "output"


# ─── Modelos ────────────────────────────────────────────────────────────────

class DocumentData(BaseModel):
    titulo: str
    nombre_empleado: str
    asunto: str
    fecha: str
    dias: str
    documento: str
    template_name: str = "memorandum.docx"


class FlujoRequest(BaseModel):
    from_date: str = "2026-01-01"
    to_date: str = "2026-04-17"
    user_ids: list[str]
    titulo: str = "Memorandum"
    template_name: str = "memorandum.docx"


# ─── Helpers ────────────────────────────────────────────────────────────────

def _replace_in_paragraph(paragraph, replacements: dict):
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(key in full_text for key in replacements):
        return
    new_text = full_text
    for placeholder, value in replacements.items():
        new_text = new_text.replace(placeholder, value)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_document(doc: Document, replacements: dict):
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)


# ─── Endpoints ──────────────────────────────────────────────────────────────

@app.post("/generar-documento", summary="Genera un documento Word manualmente")
async def generar_documento_manual(data: DocumentData):
    template_path = os.path.join(TEMPLATES_DIR, data.template_name)
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail=f"Plantilla '{data.template_name}' no encontrada.")

    replacements = {
        "<NombreEmpleado>": data.nombre_empleado,
        "<Asunto>":         data.asunto,
        "<fecha>":          data.fecha,
        "<días>":           data.dias,
        "<Documento>":      data.documento,
    }

    doc = Document(template_path)
    doc.core_properties.title = data.titulo
    _replace_in_document(doc, replacements)

    output_path = os.path.join(OUTPUT_DIR, f"{uuid.uuid4().hex}.docx")
    doc.save(output_path)

    return FileResponse(
        path=output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{data.titulo}.docx",
    )


@app.post("/flujo-automatizado", summary="Login → Obtiene datos → Genera documentos Word")
async def flujo_automatizado(request: FlujoRequest):
    # Paso 1: Login
    try:
        token = await obtener_token()
    except ValueError as e:
        raise HTTPException(status_code=401, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Error al conectar con Sapiens: {str(e)}")

    # Paso 2: Obtener registros de memorandums
    try:
        registros = await obtener_memorandums(token, request.from_date, request.to_date, request.user_ids)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Error al obtener memorandums: {str(e)}")

    if not registros:
        raise HTTPException(status_code=404, detail="No se encontraron registros para los parámetros dados.")

    # Paso 3: Generar un documento Word por cada registro
    archivos_generados = []
    for registro in registros:
        try:
            path = generar_documento(registro, titulo=request.titulo, template_name=request.template_name)
            nombre = registro.get("nombreEmpleado") or registro.get("NombreEmpleado") or "empleado"
            archivos_generados.append((path, f"{nombre}.docx"))
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error generando documento: {str(e)}")

    # Un solo archivo → devolver directo
    if len(archivos_generados) == 1:
        path, filename = archivos_generados[0]
        return FileResponse(
            path=path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename,
        )

    # Múltiples archivos → devolver ZIP
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for path, filename in archivos_generados:
            zf.write(path, arcname=filename)
    zip_buffer.seek(0)

    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename=memorandums.zip"},
    )


@app.get("/flujo-debug", summary="Muestra el token y la respuesta raw del API sin generar documentos")
async def flujo_debug(from_date: str = "2026-01-01", to_date: str = "2026-04-17"):
    try:
        token = await obtener_token()
    except Exception as e:
        return {"paso1_login": "ERROR", "error": str(e), "token": None}

    try:
        registros = await obtener_memorandums(
            token, from_date, to_date,
            ["029E5A4D-88BB-470E-1507-08DB56A77B7F"]
        )
    except Exception as e:
        return {"paso1_login": "OK", "token_preview": token[:20] + "...", "paso2_data": "ERROR", "error": str(e)}

    return {
        "paso1_login": "OK",
        "token_preview": token[:20] + "...",
        "paso2_data": "OK",
        "total_registros": len(registros),
        "primer_registro": registros[0] if registros else None,
    }


@app.get("/plantillas", summary="Lista las plantillas disponibles")
async def listar_plantillas():
    files = [f for f in os.listdir(TEMPLATES_DIR) if f.endswith(".docx")] if os.path.exists(TEMPLATES_DIR) else []
    return {"plantillas": files}


@app.get("/health")
async def health():
    return {"status": "ok"}
