"""Servicio para generar documentos Word desde plantillas."""
import os
import re
from datetime import datetime
from docx import Document

TEMPLATES_DIR = "templates"
OUTPUT_DIR = "output"


def _replace_in_paragraph(paragraph, replacements: dict[str, str]):
    """
    Reemplaza placeholders en un párrafo.
    Maneja correctamente placeholders partidos en múltiples runs por Word
    (ej: ['<', 'NombreEmpleado', '>'] → fusiona primero, luego reemplaza).
    """
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    if not any(key in full_text for key in replacements):
        return

    # Aplicar todos los reemplazos sobre el texto completo
    new_text = full_text
    for placeholder, value in replacements.items():
        new_text = new_text.replace(placeholder, str(value))

    # Copiar el formato del primer run con texto visible
    # y colapsar todos los demás runs para evitar duplicados
    first_run = paragraph.runs[0]
    first_run.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_in_document(doc: Document, replacements: dict[str, str]):
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)


def _formatear_fecha(fecha_str: str) -> str:
    """Convierte '2026-01-05T00:00:00' a '05/01/2026'."""
    if not fecha_str:
        return ""
    try:
        dt = datetime.fromisoformat(fecha_str)
        return dt.strftime("%d/%m/%Y")
    except Exception:
        return fecha_str


def _nombre_seguro(nombre: str) -> str:
    """Elimina caracteres no válidos para nombres de archivo."""
    return re.sub(r'[\\/*?:"<>|]', "", nombre).strip()


def generar_documento(registro: dict, titulo: str, template_name: str = "memorandum.docx") -> str:
    """
    Genera un .docx desde un registro del API de Sapiens.
    Retorna la ruta al archivo generado.
    """
    template_path = os.path.join(TEMPLATES_DIR, template_name)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Plantilla no encontrada: {template_path}")

    # Campos reales que devuelve la API de Sapiens
    colaborador    = registro.get("Colaborador") or registro.get("nombreEmpleado") or registro.get("NombreEmpleado") or "Empleado"
    dni            = registro.get("DNI")          or registro.get("Documento")      or registro.get("documento")      or ""
    fecha_raw      = registro.get("FechaLimite")  or registro.get("fecha")          or registro.get("Fecha")          or ""
    asunto         = registro.get("TipoIncumplimiento") or registro.get("asunto")   or registro.get("Asunto")         or ""
    cantidad_str   = str(registro.get("CantidadFaltas") or registro.get("dias")     or registro.get("Dias")           or "")
    lider          = registro.get("Lider")        or ""
    fecha_formateada = _formatear_fecha(fecha_raw)

    # En la plantilla, <días> aparece en el párrafo del cuerpo como la fecha
    # del día de incumplimiento (FechaLimite), no como cantidad de faltas.
    replacements = {
        "<NombreEmpleado>": colaborador,
        "<Asunto>":         asunto,
        "<fecha>":          fecha_formateada,
        "<días>":           fecha_formateada,   # fecha del incumplimiento
        "<Documento>":      dni,
        "<Lider>":          lider,
        # Variantes en mayúsculas
        "<NOMBREEMPLEADO>": colaborador,
        "<FECHA>":          fecha_formateada,
        "<DIAS>":           fecha_formateada,
        "<DOCUMENTO>":      dni,
    }

    doc = Document(template_path)
    doc.core_properties.title = titulo
    _replace_in_document(doc, replacements)

    # Nombre del archivo: "Nombre Completo - DD-MM-YYYY.docx" (fecha del servicio)
    fecha_archivo = fecha_formateada.replace("/", "-") if fecha_formateada else datetime.now().strftime("%d-%m-%Y")
    nombre_archivo = f"{_nombre_seguro(colaborador)} - {fecha_archivo}.docx"
    output_path = os.path.join(OUTPUT_DIR, nombre_archivo)
    doc.save(output_path)

    return output_path

